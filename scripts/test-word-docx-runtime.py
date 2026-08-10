#!/usr/bin/env python3
"""Exercise word-docx bootstrap and editing with a native, pre-downloaded wheelhouse."""

from __future__ import annotations

import argparse
import hashlib
import os
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile


def run(command: list[str], *, env: dict[str, str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        check=False,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        timeout=360,
    )


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--wheel-dir", type=Path, required=True)
    args = parser.parse_args()

    wheel_dir = args.wheel_dir.resolve(strict=True)
    wheels = list(wheel_dir.glob("*.whl"))
    if len(wheels) != 3:
        raise RuntimeError(f"expected one native lxml wheel and two universal wheels, found {len(wheels)}")

    repository = Path(__file__).resolve().parent.parent
    source = repository / "packages" / "skills" / "offline" / "word-docx"
    with tempfile.TemporaryDirectory(prefix="penguin-word-docx-native-") as temporary:
        root = Path(temporary)
        agent = root / "project" / "agents" / "default_agent"
        skill = agent / "agent_state" / "skills" / "word-docx"
        shutil.copytree(source, skill)
        offline_wheels = root / "offline" / "word-docx" / "wheels"
        offline_wheels.mkdir(parents=True)
        for wheel in wheels:
            shutil.copy2(wheel, offline_wheels / wheel.name)

        environment = os.environ.copy()
        for key in tuple(environment):
            if key.startswith("PIP_") or key in {"PYTHONHOME", "PYTHONPATH"}:
                environment.pop(key)
        escaped = root / "escaped-pip-target"
        environment.update(
            {
                "PENGUIN_OFFLINE_ROOT": str(root / "offline"),
                "PIP_TARGET": str(escaped),
                "PIP_USER": "1",
                "PYTHONHOME": str(root / "invalid-python-home"),
                "PYTHONPATH": str(root / "invalid-python-path"),
            }
        )
        bootstrap = skill / "scripts" / "bootstrap.py"

        # Two simultaneous first calls exercise the platform lock and must converge on one venv.
        processes = [
            subprocess.Popen(
                [sys.executable, "-I", str(bootstrap), "--help"],
                env=environment,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
            )
            for _ in range(2)
        ]
        for process in processes:
            output, _ = process.communicate(timeout=360)
            if process.returncode != 0:
                raise RuntimeError(f"concurrent bootstrap failed ({process.returncode}):\n{output}")
        if escaped.exists():
            raise RuntimeError("pip environment escaped the Agent-owned virtual environment")

        python = (
            agent / "shared_env" / "word-docx" / "Scripts" / "python.exe"
            if os.name == "nt"
            else agent / "shared_env" / "word-docx" / "bin" / "python"
        )
        imported = run([str(python), "-I", "-c", "import docx, lxml"], env=environment)
        if imported.returncode != 0:
            raise RuntimeError(f"controlled environment imports failed:\n{imported.stdout}")

        input_path = root / "input-without-heading.docx"
        output_path = root / "output.docx"
        create = run(
            [
                str(python),
                "-I",
                "-c",
                (
                    "from docx import Document; "
                    "d=Document(); h=d.styles['Heading 1']; d.styles.element.remove(h._element); "
                    "d.add_paragraph('Existing paragraph'); d.save(r'%s')"
                )
                % input_path,
            ],
            env=environment,
        )
        if create.returncode != 0:
            raise RuntimeError(f"DOCX fixture creation failed:\n{create.stdout}")
        source_digest = digest(input_path)

        appended = run(
            [
                sys.executable,
                "-I",
                str(bootstrap),
                "append",
                "--input",
                str(input_path),
                "--output",
                str(output_path),
                "--heading",
                "Native offline heading",
                "--paragraph",
                "Native offline paragraph",
            ],
            env=environment,
        )
        if appended.returncode != 0:
            raise RuntimeError(f"DOCX append failed:\n{appended.stdout}")
        if digest(input_path) != source_digest:
            raise RuntimeError("source DOCX changed")

        verify = run(
            [
                str(python),
                "-I",
                "-c",
                (
                    "from docx import Document; d=Document(r'%s'); "
                    "texts=[p.text for p in d.paragraphs]; "
                    "assert 'Existing paragraph' in texts; "
                    "assert 'Native offline heading' in texts; "
                    "assert 'Native offline paragraph' in texts; "
                    "p=next(p for p in d.paragraphs if p.text=='Native offline heading'); "
                    "assert any(r.bold is True for r in p.runs)"
                )
                % output_path,
            ],
            env=environment,
        )
        if verify.returncode != 0:
            raise RuntimeError(f"DOCX reopen verification failed:\n{verify.stdout}")

    print(f"word-docx native runtime acceptance passed: {sys.platform}/{os.uname().machine if hasattr(os, 'uname') else os.environ.get('PROCESSOR_ARCHITECTURE', 'unknown')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
