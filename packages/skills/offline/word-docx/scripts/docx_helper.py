#!/usr/bin/env python3
"""Deterministic, non-destructive DOCX inspection and basic editing commands."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
from pathlib import Path
import tempfile
from typing import Callable

from docx import Document
from docx.oxml.ns import qn


ORDINARY_RUN_CHILDREN = {qn(name) for name in ("w:rPr", "w:t", "w:tab", "w:br", "w:cr")}


def file_digest(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def input_path(raw: str) -> Path:
    path = Path(raw).expanduser().resolve(strict=True)
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise ValueError(f"input must be an existing .docx file: {path}")
    return path


def output_path(raw: str, source: Path) -> Path:
    path = Path(raw).expanduser().resolve()
    if path.suffix.lower() != ".docx":
        raise ValueError(f"output must use the .docx extension: {path}")
    if path == source:
        raise ValueError("output must be different from input")
    if path.exists():
        raise FileExistsError(f"output already exists: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def paragraph_texts(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def is_ordinary_text_run(run) -> bool:
    return all(child.tag in ORDINARY_RUN_CHILDREN for child in run._r)


def save_verified(
    document: Document,
    source: Path,
    output: Path,
    verify: Callable[[Document], None],
) -> None:
    source_before = file_digest(source)
    temporary: Path | None = None
    linked = False
    try:
        with tempfile.NamedTemporaryFile(
            prefix=f".{output.name}.", suffix=".tmp", dir=output.parent, delete=False
        ) as handle:
            temporary = Path(handle.name)
        document.save(temporary)
        verify(Document(temporary))
        os.link(temporary, output)
        linked = True
        verify(Document(output))
        if file_digest(source) != source_before:
            raise RuntimeError("source DOCX changed during editing")
    except Exception:
        if linked:
            output.unlink(missing_ok=True)
        raise
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def inspect_document(args: argparse.Namespace) -> dict[str, object]:
    source = input_path(args.input)
    document = Document(source)
    paragraphs = [
        {"index": index, "style": paragraph.style.name, "text": paragraph.text}
        for index, paragraph in enumerate(document.paragraphs)
    ]
    return {
        "operation": "inspect",
        "input": str(source),
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "verified": True,
    }


def append_heading(document: Document, text: str, level: int) -> None:
    style_name = f"Heading {level}"
    if style_name in document.styles:
        document.add_heading(text, level=level)
        return

    # Some third-party DOCX files omit Word's built-in heading styles entirely.
    # Keep the edit usable without mutating the source or inventing a document-wide style.
    run = document.add_paragraph().add_run(text)
    run.bold = True


def append_content(args: argparse.Namespace) -> dict[str, object]:
    if args.heading is None and not args.paragraph:
        raise ValueError("append requires --heading or at least one --paragraph")
    source = input_path(args.input)
    output = output_path(args.output, source)
    document = Document(source)
    expected: list[str] = []
    if args.heading is not None:
        append_heading(document, args.heading, args.heading_level)
        expected.append(args.heading)
    for text in args.paragraph:
        document.add_paragraph(text)
        expected.append(text)

    def verify(reopened: Document) -> None:
        texts = paragraph_texts(reopened)
        if texts[-len(expected) :] != expected:
            raise RuntimeError("appended content is missing after reopen")

    save_verified(document, source, output, verify)
    return {
        "operation": "append",
        "input": str(source),
        "output": str(output),
        "appended": expected,
        "verified": True,
    }


def replace_text(args: argparse.Namespace) -> dict[str, object]:
    if args.old == "":
        raise ValueError("--old must not be empty")
    source = input_path(args.input)
    output = output_path(args.output, source)
    document = Document(source)
    matching_runs = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            count = run.text.count(args.old)
            if count:
                if not is_ordinary_text_run(run):
                    raise ValueError(
                        "text occurs in a run containing a drawing, object, field, or other "
                        "non-text content; refusing a destructive replacement"
                    )
                matching_runs.append((run, count))
    if not matching_runs:
        raise ValueError(f"text was not found within an ordinary text run: {args.old!r}")
    replacements = sum(count for _, count in matching_runs)
    for run, _ in matching_runs:
        run.text = run.text.replace(args.old, args.new)
    expected_texts = paragraph_texts(document)

    def verify(reopened: Document) -> None:
        if paragraph_texts(reopened) != expected_texts:
            raise RuntimeError("replacement content differs after reopen")

    save_verified(document, source, output, verify)
    return {
        "operation": "replace",
        "input": str(source),
        "output": str(output),
        "replacements": replacements,
        "verified": True,
    }


def parser() -> argparse.ArgumentParser:
    root = argparse.ArgumentParser(description=__doc__)
    commands = root.add_subparsers(dest="command", required=True)

    inspect = commands.add_parser("inspect", help="read basic paragraphs")
    inspect.add_argument("--input", required=True)
    inspect.set_defaults(run=inspect_document)

    append = commands.add_parser("append", help="append a heading and/or paragraphs")
    append.add_argument("--input", required=True)
    append.add_argument("--output", required=True)
    append.add_argument("--heading")
    append.add_argument("--heading-level", type=int, choices=range(1, 10), default=1)
    append.add_argument("--paragraph", action="append", default=[])
    append.set_defaults(run=append_content)

    replace = commands.add_parser("replace", help="replace text in ordinary paragraphs")
    replace.add_argument("--input", required=True)
    replace.add_argument("--output", required=True)
    replace.add_argument("--old", required=True)
    replace.add_argument("--new", required=True)
    replace.set_defaults(run=replace_text)
    return root


def main() -> int:
    args = parser().parse_args()
    try:
        result = args.run(args)
    except Exception as error:
        print(json.dumps({"error": str(error), "verified": False}, ensure_ascii=False))
        return 1
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
